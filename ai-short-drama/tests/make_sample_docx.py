"""生成一个示例短剧 .docx 用于端到端验证."""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document


def make_sample(output_path: Path) -> None:
    doc = Document()
    doc.add_heading("逆袭:被冤枉的设计师", level=0)

    doc.add_heading("第一场 办公室对峙", level=1)
    doc.add_paragraph("林晓站在主管面前，握紧手中的辞退通知单。")
    doc.add_paragraph("主管陈伟冷笑：'这次抄袭事件，公司决定让你承担全部责任。'")
    doc.add_paragraph("林晓抬眼，目光坚定：'方案是我熬夜三天做的，不是抄袭。'")
    doc.add_paragraph("陈伟拍桌：'证据不会说谎。明天交接，下周离职。'")

    doc.add_heading("第二场 深夜独白", level=1)
    doc.add_paragraph("林晓独自坐在办公室，电脑屏幕的光打在脸上。")
    doc.add_paragraph("她调出三天前的版本历史，发现陈伟的助理在后台修改了源文件时间戳。")
    doc.add_paragraph("一封匿名邮件弹出：'有人想让你背锅，证据在我手里。'")

    doc.add_heading("第三场 真相浮出", level=1)
    doc.add_paragraph("会议室里，林晓打开投影仪，调出 Git 提交记录。")
    doc.add_paragraph("'各位，每一行修改都有时间戳和作者签名。'")
    doc.add_paragraph("陈伟脸色铁青：'这……这是伪造的。'")
    doc.add_paragraph("CEO 走进来：'伪造？我们的安全审计可不会作假。'")

    doc.add_heading("第四场 反转", level=1)
    doc.add_paragraph("CEO 拍着林晓的肩：'设计部主管位置空了，你来。'")
    doc.add_paragraph("陈伟被两个保安押出会议室。")
    doc.add_paragraph("林晓望向窗外，城市的灯火在玻璃上闪烁。")
    doc.add_paragraph("她拿起手机，给母亲发消息：'妈，我升职了。'")

    doc.save(output_path)
    print(f"[OK] 示例剧本已生成 → {output_path}")


if __name__ == "__main__":
    out = Path(__file__).resolve().parent.parent / "tests" / "sample_script.docx"
    make_sample(out)
    print(f"\n用法:")
    print(f"  streamlit run app.py")
    print(f"  上传 {out}")
