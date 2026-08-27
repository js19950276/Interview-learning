from __future__ import annotations

import io
import json
import logging

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from agent.storyboard import Storyboard

log = logging.getLogger("exporter")


def export_rewrite(title: str, emotion_type: str, text: str) -> io.BytesIO:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(12)

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"情绪类型：{emotion_type}")
    doc.add_paragraph("")

    for line in text.split("\n"):
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    log.info("重写文档导出完成 | 标题=%s", title)
    return buf


def export_storyboard_word(title: str, storyboard: Storyboard) -> io.BytesIO:
    """分镜表格导出为 Word."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(12)

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"情绪类型：{storyboard.emotion_type}")
    doc.add_paragraph(f"总时长：{storyboard.total_duration_seconds} 秒")
    doc.add_paragraph(f"镜头数：{len(storyboard.shots)}")
    doc.add_paragraph("")

    headers = ["#", "片段编号", "源段", "景别", "时长(秒)", "画面描述", "细节"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr[i].text = h_text

    for shot in storyboard.shots:
        row = table.add_row().cells
        row[0].text = str(shot.shot_number)
        row[1].text = shot.fragment_id or "-"
        row[2].text = shot.source_anchor or "-"
        row[3].text = shot.shot_type
        row[4].text = str(shot.duration_seconds)
        row[5].text = shot.visual_description
        details = []
        if shot.dialogue:
            details.append(f"台词: {shot.dialogue}")
        if shot.sound_effects:
            details.append(f"音效: {shot.sound_effects}")
        if shot.bgm_mood:
            details.append(f"BGM: {shot.bgm_mood}")
        if shot.emotion_note:
            details.append(f"情绪: {shot.emotion_note}")
        if shot.camera_movement:
            details.append(f"运镜: {shot.camera_movement}")
        row[6].text = "\n".join(details) if details else "-"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    log.info("分镜 Word 导出完成 | 标题=%s | 镜头=%d", title, len(storyboard.shots))
    return buf


def export_video_prompts(storyboard: Storyboard) -> io.BytesIO:
    """导出 video_prompts.json — 直接喂 AI 视频模型(全中文 prompt)."""
    payload = {
        "title": storyboard.title,
        "emotion_type": storyboard.emotion_type,
        "total_duration_seconds": storyboard.total_duration_seconds,
        "fragments": [
            {
                "fragment_id": shot.fragment_id,
                "source_anchor": shot.source_anchor,
                "duration": shot.duration_seconds,
                "shot_type": shot.shot_type,
                "prompt": shot.visual_description,
                "negative_prompt": shot.negative_prompt,
                "camera_movement": shot.camera_movement,
                "audio": {
                    "dialogue": shot.dialogue,
                    "sound_effects": shot.sound_effects,
                    "bgm_mood": shot.bgm_mood,
                },
            }
            for shot in storyboard.shots
        ],
    }
    buf = io.BytesIO(json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8"))
    buf.seek(0)
    log.info("video_prompts.json 导出完成 | 镜头=%d", len(storyboard.shots))
    return buf
